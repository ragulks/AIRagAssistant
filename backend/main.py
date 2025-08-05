import ollama
from typing import List, Tuple, Optional
import time
import os
from pathlib import Path
import re

# For document processing
from docx import Document
import PyPDF2
import json

# Configuration
EMBEDDING_MODEL = 'nomic-embed-text'  # Good local embedding model
LANGUAGE_MODEL = 'llama3'  # Default local Llama3 model

# Each element in the VECTOR_DB will be a tuple (chunk, embedding, source)
VECTOR_DB: List[Tuple[str, List[float], str]] = []


class DocumentProcessor:
    """Handle different document types and text processing."""

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks."""
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to end at a sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap

        return chunks

    @staticmethod
    def load_txt_file(file_path: str) -> str:
        """Load text from a TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    @staticmethod
    def load_docx_file(file_path: str) -> str:
        """Load text from a DOCX file."""
        doc = Document(file_path)
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))

        return '\n'.join(text_content)

    @staticmethod
    def load_pdf_file(file_path: str) -> str:
        """Load text from a PDF file."""
        text_content = []

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_content.append(text.strip())

        return '\n'.join(text_content)

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text."""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-"]', '', text)
        return text.strip()


def load_document(file_path: str) -> List[str]:
    """Load and process a document into chunks."""
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    processor = DocumentProcessor()

    # Determine file type and load content
    if file_path.suffix.lower() == '.txt':
        raw_text = processor.load_txt_file(file_path)
    elif file_path.suffix.lower() == '.docx':
        raw_text = processor.load_docx_file(file_path)
    elif file_path.suffix.lower() == '.pdf':
        raw_text = processor.load_pdf_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")

    # Clean the text
    clean_text = processor.clean_text(raw_text)

    # Split into chunks
    chunks = processor.chunk_text(clean_text)

    return chunks


def add_chunk_to_database(chunk: str, source: str) -> None:
    """Add a text chunk to the vector database with its embedding."""
    try:
        # Generate embedding for the chunk
        response = ollama.embeddings(
            model=EMBEDDING_MODEL,
            prompt=chunk
        )
        embedding = response['embedding']
        VECTOR_DB.append((chunk, embedding, source))
    except Exception as e:
        print(f"Error processing chunk: {e}")


def clear_database():
    """Clear the vector database."""
    global VECTOR_DB
    VECTOR_DB = []


def cosine_similarity(a: List[float], b: List[float]) -> float:
    """Calculate cosine similarity between two vectors."""
    dot_product = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x ** 2 for x in a) ** 0.5
    norm_b = sum(x ** 2 for x in b) ** 0.5
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot_product / (norm_a * norm_b)


def retrieve(query: str, top_n: int = 5) -> List[Tuple[str, float, str]]:
    """Retrieve top_n most similar chunks to the query."""
    try:
        # Get embedding for the query
        response = ollama.embeddings(
            model=EMBEDDING_MODEL,
            prompt=query
        )
        query_embedding = response['embedding']

        # Calculate similarities
        similarities = []
        for chunk, embedding, source in VECTOR_DB:
            similarity = cosine_similarity(query_embedding, embedding)
            similarities.append((chunk, similarity, source))

        # Sort by similarity and return top_n
        similarities.sort(key=lambda x: x[1], reverse=True)
        return similarities[:top_n]
    except Exception as e:
        print(f"Error during retrieval: {e}")
        return []


def generate_response(query: str, retrieved_chunks: List[Tuple[str, float, str]]) -> str:
    """Generate response using the language model."""
    if not retrieved_chunks:
        return "I don't have enough information to answer that question."

    # Build context from retrieved chunks
    context = '\n'.join([f'- {chunk}' for chunk, _, _ in retrieved_chunks])

    prompt = f"""You are a helpful AI assistant. Answer the question based only on the following context information:

Context:
{context}

Question: {query}

Instructions:
- Answer based only on the provided context
- If the context doesn't contain relevant information, say "I don't have enough information to answer that question based on the provided documents."
- Be concise and accurate
- Use Portuguese if the question is in Portuguese, otherwise use English

Answer:"""

    try:
        # Generate response (non-streaming for API compatibility)
        response = ollama.generate(
            model=LANGUAGE_MODEL,
            prompt=prompt,
            stream=False
        )
        return response['response']
    except Exception as e:
        return f"Error generating response: {e}"


def process_uploaded_file(file_path: str) -> bool:
    """Process an uploaded file and add it to the database."""
    try:
        print(f"Processing file: {file_path}")

        # Clear existing database
        clear_database()

        # Load and process the document
        chunks = load_document(file_path)
        print(f"Document split into {len(chunks)} chunks")

        # Add chunks to database
        for i, chunk in enumerate(chunks, 1):
            add_chunk_to_database(chunk, os.path.basename(file_path))
            if i % 10 == 0 or i == len(chunks):
                print(f'Processed {i}/{len(chunks)} chunks')

        print(f"Successfully processed {len(chunks)} chunks from {file_path}")
        return True

    except Exception as e:
        print(f"Error processing file {file_path}: {e}")
        return False


def chat_query(query: str) -> dict:
    """Process a chat query and return response with metadata."""
    if not VECTOR_DB:
        return {
            "response": "No documents have been uploaded yet. Please upload a document first.",
            "sources": [],
            "chunks_used": 0
        }

    # Retrieve relevant chunks
    retrieved_chunks = retrieve(query)

    if not retrieved_chunks:
        return {
            "response": "I couldn't find any relevant information in the uploaded documents.",
            "sources": [],
            "chunks_used": 0
        }

    # Generate response
    response = generate_response(query, retrieved_chunks)

    # Extract unique sources
    sources = list(set([source for _, _, source in retrieved_chunks]))

    return {
        "response": response,
        "sources": sources,
        "chunks_used": len(retrieved_chunks)
    }


def main():
    """Main function for command line usage."""
    print("Enhanced RAG Document Chatbot")
    print("Supports: TXT, DOCX, PDF files")
    print("-" * 50)

    # Load initial document if specified
    default_file = input("Enter path to initial document (or press Enter to skip): ").strip()

    if default_file and os.path.exists(default_file):
        if process_uploaded_file(default_file):
            print("✓ Document loaded successfully!")
        else:
            print("✗ Failed to load document")

    print("\nChatbot ready!")
    print("Commands:")
    print("- 'upload <filepath>' to load a new document")
    print("- 'quit' to exit")
    print("-" * 50)

    while True:
        try:
            user_input = input('\nYou: ').strip()

            if user_input.lower() in ['quit', 'exit', 'q']:
                break

            if user_input.lower().startswith('upload '):
                file_path = user_input[7:].strip()
                if process_uploaded_file(file_path):
                    print("✓ Document uploaded and processed successfully!")
                else:
                    print("✗ Failed to process document")
                continue

            if not user_input:
                continue

            # Process query
            result = chat_query(user_input)

            print(f"\nChatbot: {result['response']}")

            if result['sources']:
                print(f"\nSources: {', '.join(result['sources'])}")
                print(f"Chunks used: {result['chunks_used']}")

        except KeyboardInterrupt:
            print("\nGoodbye!")
            break
        except Exception as e:
            print(f"\nAn error occurred: {e}")


if __name__ == "__main__":
    main()